import os
import collections
import re
import colorama
import pdb
import prettytable
from time import sleep
from docx import Document
from docx.shared import Inches, Cm, RGBColor
from docx2pdf import convert
from docx.enum.text import WD_ALIGN_PARAGRAPH
from guessit import guessit


class File():
    '''Carries info of a file and has file management methods'''

    def __init__(self, path, file_type, dir_tag='', size=0, surface=False):
        self.name = os.path.basename(path)
        self.path = path
        self.dir_tag = dir_tag
        self.file_type = file_type
        self.size = size
        self.surface = surface

    def get_size(self):
        '''returns size(int or float) of file and unit(str) of that size'''
        # size is in bytes
        # 10^3 > KB
        # 10^6 > MB
        if self.size < 10**6:  # converts bytes to KBs
            size = self.size/1024
            unit = 'KB'
            size = int(size)
        elif self.size < 10**9:  # converts bytes to MBs
            # 1048576 is 1024 * 1024
            size = self.size/1048576
            size = round(size, 1)  # rounds to a total of 4 digits
            unit = 'MB'
        else:  # if size>10^9 converts to GBs
            # 1073741824 is 1024 * 1024 * 1024
            size = self.size/1073741824
            size = round(size, 1)  # rounds to a total of 4 digits
            unit = 'GB'
        if int(size) == 0:
            size = 0
        return size, unit

    def set_index(self, index):
        self.index = index


def parse_inp(inp):
    net_inp, cmd = get_cmd(inp)
    net_inp, dir_tag = get_dir_tag(inp)
    if cmd:
        cmd['type'] = 'cmd'
    elif dir_tag:
        cmd['type'] = 'dir tag'
        cmd['dir tag'] = dir_tag
    return net_inp, cmd


def get_cmd(inp):
    '''Returns a dictionary containing the details of the command written'''
    match = re.search('-.*', inp)
    cmd = {}
    if match:
        cmd_tag = match.group()
        net_inp = inp.replace(cmd_tag, '')
    else:
        return inp, cmd

    if cmd_tag == '-dups':  # to change what function catches as commands
        # just change the if coniditon nothing else
        cmd['func name'] = 'print_duplicates'
    elif cmd_tag == '-all movies':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'movie'
    elif cmd_tag == '-all tv':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'tv series'
    elif cmd_tag == '-all animes':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'anime'
    elif cmd_tag == '-this pc e':
        cmd['func name'] = 'get_clips'
        cmd['dir tag'] = 'this pc e'
    elif cmd_tag == '-this pc f':
        cmd['func name'] = 'get_clips'
        cmd['dir tag'] = 'this pc f'
    elif cmd_tag == '-this pc tv':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'tv series'
        cmd['dir tag'] = 'this pc'
    elif cmd_tag == '-this pc movies':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'movie'
        cmd['dir tag'] = 'this pc'
    elif cmd_tag == '-2 tera movies':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'movie'
        cmd['dir tag'] = '2 tera movies'
    elif cmd_tag == '-2 tera tv':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'tv series'
        cmd['dir tag'] = '2 tera tv'
    elif cmd_tag == '-2 tera animes':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'anime'
        cmd['dir tag'] = '2 tera animes'
    elif cmd_tag == '-1 tera movies':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'movie'
        cmd['dir tag'] = '1 tera movies'
    elif cmd_tag == '-1 tera tv':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'tv series'
        cmd['dir tag'] = '1 tera tv'
    elif cmd_tag == '-1 tera animes':
        cmd['func name'] = 'get_clips'
        cmd['file type'] = 'anime'
        cmd['dir tag'] = '1 tera animes'
    return net_inp, cmd


def get_dir_tag(inp):
    '''returns directory tag(if exists) and input line without the tag in it'''
    match = re.search(r'''^         ((this \s pc (\s(e|f))? (?=\s))  
                        # this pc followed by e or f opt.,followed by space(not consumed)
                                                        |
                                ((1|2) \s tera (\s (movies|tv|animes) )? (?=\s)))''',
                      # 1|2 tera followed by movies tv or animes(opt.) followed by space(not consumed)
                      inp,
                      re.VERBOSE | re.I)
    if match:
        dir_tag = match.group()
        to_replace = dir_tag+' '  # because dir_tag doesn't have the space in it
        net_inp = inp.replace(to_replace, '')
    else:
        dir_tag = None
        net_inp = inp
    return net_inp, dir_tag


def search(paths, search_for, dir_tag=''):
    '''returns a list of File class objects thatcarry the results's data/details.'''
    os.system('cls')
    search_for = search_for.lower().split()
    SEARCH_WORDS_COUNT = len(search_for)
    results = []
    index = 0
    for path in paths:
        matches_count = 0
        if path.surface and search_for[0] in path.name.lower():
            for word in search_for:
                if word in path.name.lower():
                    matches_count += 1
            if matches_count == SEARCH_WORDS_COUNT:
                if dir_tag:  # i think if you remove this
                    if dir_tag in path.dir_tag.lower():
                        path.set_index(index)
                        results.append(path)
                        index += 1
                else:  # and this the code will work thus fine because '' is in any string anyway
                    path.set_index(index)
                    results.append(path)
                    index += 1
    return results


def print_results(results, dups=False):
    '''
    takes in results of a search (File objects or Dup objects) and prints their info in a formatted
    manner
     '''
    colorama.init()
    os.system('cls')
    print(colorama.Fore.YELLOW+'\n\n\t\t\t\t\t\t\tResults\n\n\n')
    if not dups:
        if len(results) != 0:
            table = prettytable.PrettyTable([colorama.Fore.YELLOW+'N',
                                                colorama.Fore.YELLOW+'Names',
                                                colorama.Fore.YELLOW+'Dir Tag',
                                                colorama.Fore.YELLOW+'Size'])
            num = 1
            for result in results:
                size, unit = result.get_size()
                table.add_row([colorama.Fore.LIGHTBLUE_EX+f'{num}',
                                colorama.Fore.LIGHTBLUE_EX+f'{result.name:80}',
                                colorama.Fore.LIGHTBLUE_EX+f'{result.dir_tag:10}',
                                colorama.Fore.LIGHTBLUE_EX+f'{size:<6} {unit}'])
                num += 1
                # print(colorama.Fore.LIGHTBLUE_EX + f'{result.name:75}             {result.dir_tag:50}\n')
            print(table)
        else:
            print(' ' * 51 + '----No results----')
    else:
        for result in results:
            if len(result.dir_tags) > 4:  # prints only 3 and .... the rest
                dir_tags = results.dir_tags[:3]
                dir_tags.append('....')
            else:
                dir_tags = result.dir_tags
            dir_tags = ' | '.join(dir_tags)
            print(colorama.Fore.RED + f'{result.name:70} {dir_tags:50}\n')


def get_duplicates(clips):
    '''
    returns a list of Dup namedtuples that contain the info of the duplicate
    surface files that exist in trees
    '''
    print('Guessing names of files..')
    clips_names = [guessit(clip.name)['title']
                   for clip in clips if clip.surface]
    files_count = collections.Counter(clips_names)
    duplicates = []
    dups_names = []
    for file_name, count in files_count.items():
        if count > 1:
            dups_names.append(file_name)
    for dup_name in dups_names:
        results = search(clips, dup_name)  # all results are of same name
        # but they have diff dir_tags/locations
        locations = [result.path for result in results]
        dir_tags = [result.dir_tag for result in results]
        file_type = results[0].file_type
        Dup = collections.namedtuple(
            'dup', 'name file_type locations dir_tags')
        dup_obj = Dup(dup_name, file_type, locations, dir_tags)
        duplicates.append(dup_obj)
    return duplicates


def to_unit(size):
    # size is in bytes
    # 10^3 > KB
    # 10^6 > MB
    if size < 10**6:  # converts bytes to KBs
        size = size/1024
        unit = 'KB'
    elif size < 10**9:  # converts bytes to MBs
        # 1048576 is 1024 * 1024
        size = size/1048576
        unit = 'MB'
    else:  # if size>10^9 converts to GBs
        # 1073741824 is 1024 * 1024 * 1024
        size = size/1073741824
        unit = 'GB'
    size = round(size, 4)  # rounds to a totale of 4 digits
    size = str(size)
    return size, unit


def print_tree(paths):
    '''prints tree with indentation that indicate which folder/file is inside which'''
    os.system('cls')
    cont_path = paths[0]
    print(f'{colorama.Fore.YELLOW}\n\n\t\t\t\t\t\t\tResults\n\n\n')
    cont_size, cont_unit = cont_path.get_size()
    print(colorama.Fore.YELLOW +
          f'\n\n{cont_path.name:70}', end='')  # has no indentation
    print(colorama.Fore.RED +
          f'{cont_size:12} {cont_unit}\n')
    print(colorama.Fore.YELLOW + '-' * 90 + '\n')
    for path in paths[1:]:
        # you can simply know the level by the number of backslashes in the name
        # we remove the backslashes in start to restore the level to 0 for the levels
        # to be in respect to the start_path itself not C:\ or F:\
        level = path.path.replace(cont_path.path, '').count(os.sep) - 1
        indent = ' ' * 8 * level
        size, unit = path.get_size()
        print(colorama.Fore.WHITE +
              f'{indent}{path.name:70}   {colorama.Fore.BLUE}{size:>6} {unit:>}\n')
    input()


def get_clips(paths, cmd):
    '''
    returns a chunk of clips according to dir tag and file type carried by the Cmd Object only(ex:-this pc movies)
    '''
    clips = []
    index = 0
    if len(cmd.keys()) == 3:
        if 'dir tag' in cmd.keys():
            for path in paths:
                if path.surface and cmd['dir tag'] == path.dir_tag.lower():
                    path.set_index(index)
                    clips.append(path)
                    index += 1
        elif 'file type' in cmd.keys():
            for path in paths:
                if path.surface and cmd['file type'] == path.file_type:
                    path.set_index(index)
                    clips.append(path)
                    index += 1
    elif 'dir tag' in cmd.keys():  # it means a a specific dir is to be searched
        for path in paths:
            if path.surface and cmd['dir tag'] in path.dir_tag.lower() and path.file_type == cmd['file type']:
                path.set_index(index)
                clips.append(path)
                index += 1
    return clips


def get_tree(paths, results, end_cmd):
    '''
    returns a list of the tree of the current result's file path
                    (actual full path is used to search)
    '''
    match = re.search(r'tree \d{1,}', end_cmd)
    if match:
        num = match.group().split(' ')[1]
        inp_num = int(num)
    else:
        print('something wrong with the cmd..')
        sleep(1)
    if inp_num > len(results):
        print("out of index")
    for result in results:
        if result.index + 1 == inp_num:
            root = result
            break
    tree = []
    # because result.path isn't in in it's dirname, look down
    tree.append(root)
    for path in paths:
        if root.path in os.path.dirname(path.path):
            tree.append(path)
    return tree


def clean():
    '''Cleans the folder that contains the program from .docx and .pdf files'''
    count = 0
    for f in os.listdir(os.getcwd()):
        if '.pdf' in f or 'docx' in f:
            os.remove(f)
            count += 1
    print(f'Found {count} files.')
    sleep(2)


def create_pdf(results, inp, pdf_name='results',):
    '''Creates a pdf file from results that has the name passed to function '''
    def create_table(docx, results):
        '''creates a table in the docx file object and fills the info from
        results'''
        table = docx.add_table(rows=len(results)+1, cols=4)
        row = table.rows[0]
        row.cells[0].width = Inches(13)
        row.cells[0].text = 'Name'
        row.cells[1].width = Inches(0.1)
        row.cells[1].text = 'Size'
        row.cells[2].width = Inches(0.5)
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        row.cells[2].text = ''
        row.cells[3].text = 'Location'
        for index, result in enumerate(results):
            row = table.rows[index+1]
            row.cells[0].width = Inches(11)
            row.cells[0].text = result.name
            row.cells[1].width = Inches(0.1)
            size, unit = result.get_size()
            size = str(size)
            row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row.cells[1].text = f'{size:<3}'
            row.cells[2].width = Inches(0.5)
            row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            row.cells[2].text = f'{unit:<3}'
            row.cells[3].width = Inches(4)
            row.cells[3].text = result.dir_tag

    def create_docx(results, dest_path, docx_name='temp'):
        '''Takes in a list of results and the word file name(without .docx)
        and returns path of the created file'''
        print('Creating Docx...')
        docx_name += '.docx'
        docx = Document()
        for sec in docx.sections:
            sec.top_margin = Cm(0.5)
            sec.left_margin = Cm(0.01)
            sec.bottom_margin = Cm(0.5)
            sec.right_margin = Cm(0.5)
        print('Creating table.')
        create_table(docx, results)
        print('table created...')
        if dest_path:
            DOCX_PATH = f'{dest_path}{os.path.sep}{docx_name}'
        docx.save(DOCX_PATH)
        print('Docx file created succesfuly.')
        return DOCX_PATH

    match = re.search('export to .*', inp)
    if match:
        location = match.group().split(' to ')[1]
        if location == 'desk':
            dest_path = r'C:\Users\COMPU1\Desktop'
        else:
            dest_path = ''
    else:
        dest_path = os.getcwd()

    DOCX_PATH = create_docx(results, dest_path, pdf_name)
    print(f'Converting {DOCX_PATH} to PDF.')
    convert(DOCX_PATH)
    os.remove(DOCX_PATH)
    print(' .docx Converted succesfully.')
    sleep(2)
